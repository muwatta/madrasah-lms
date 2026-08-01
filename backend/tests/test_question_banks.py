import io
import pytest
from datetime import date
from django.core.files.uploadedfile import SimpleUploadedFile
from rest_framework import status

from curriculum.models import SchoolClass
from academic.models import Session, Term
from quizzes.models import Question as QuizQuestion
from question_banks.models import QuestionBank


def make_docx():
    from docx import Document
    doc = Document()
    doc.add_paragraph('Tajweed Examination')
    doc.add_paragraph('1. What is the meaning of madd?')
    doc.add_paragraph('A. Lengthening')
    doc.add_paragraph('B. Shortening')
    doc.add_paragraph('C. Stopping')
    doc.add_paragraph('D. Pausing')
    doc.add_paragraph('Answer: A')
    doc.add_paragraph('2. Explain the rule of idgham.')
    doc.add_paragraph('3. Qalqalah letters are:')
    doc.add_paragraph('A. ق ط ب ج د')
    doc.add_paragraph('B. أ ه ع ح غ')
    doc.add_paragraph('Answer: A')
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@pytest.fixture
def school_class(madrasah):
    return SchoolClass.objects.create(
        madrasah=madrasah, name_en='Class 1', name_ar='الصف الأول', order=1)


@pytest.fixture
def session(madrasah):
    return Session.objects.create(
        madrasah=madrasah, name='2025/2026',
        start_date=date(2025, 9, 1), end_date=date(2026, 7, 31))


@pytest.fixture
def term(madrasah, session):
    return Term.objects.create(
        madrasah=madrasah, session=session, name='Term 1',
        term_number=1, start_date=date(2025, 9, 1), end_date=date(2025, 12, 20))


@pytest.fixture
def bank(madrasah, teacher, subject, school_class, session, term):
    return QuestionBank.objects.create(
        madrasah=madrasah, created_by=teacher, subject=subject,
        school_class=school_class, session=session, term=term,
        title='Tajweed Term 1', file_type='docx', status='ready')


@pytest.mark.django_db
def test_teacher_uploads_docx_and_questions_parsed(
        teacher_client, madrasah, subject, school_class, session, term):
    docx = make_docx()
    response = teacher_client.post(
        '/api/v1/question-banks/',
        {
            'file': SimpleUploadedFile('exam.docx', docx.read(),
                                      content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
            'subject': subject.id,
            'school_class': school_class.id,
            'session': session.id,
            'term': term.id,
            'title': 'Tajweed Term 1 Exam',
        },
        format='multipart',
    )
    assert response.status_code == status.HTTP_201_CREATED, response.content
    data = response.data
    assert data['status'] == 'ready'
    assert data['question_count'] == 3
    assert data['file_type'] == 'docx'
    assert data['size_saved'] >= 0

    questions = list(QuestionBank.objects.get(pk=data['id']).bank_questions.order_by('id'))
    assert len(questions) == 3
    mcq = questions[0]
    assert mcq.question_type == 'mcq'
    assert mcq.correct_answer == 'A'
    assert len(mcq.options) == 4
    sa = questions[1]
    assert sa.question_type == 'short_answer'
    assert sa.correct_answer == ''
    assert questions[2].question_type == 'mcq'


@pytest.mark.django_db
def test_only_one_bank_per_term(
        teacher_client, madrasah, subject, school_class, session, term, bank):
    docx = make_docx()
    response = teacher_client.post(
        '/api/v1/question-banks/',
        {
            'file': SimpleUploadedFile('exam.docx', docx.read(), content_type='application/octet-stream'),
            'subject': subject.id,
            'school_class': school_class.id,
            'session': session.id,
            'term': term.id,
        },
        format='multipart',
    )
    assert response.status_code == status.HTTP_400_BAD_REQUEST
    assert 'already uploaded' in response.data['error']


@pytest.mark.django_db
def test_invalid_file_type_rejected(teacher_client, madrasah, subject, school_class, session, term):
    response = teacher_client.post(
        '/api/v1/question-banks/',
        {
            'file': SimpleUploadedFile('exam.txt', b'hello', content_type='text/plain'),
            'subject': subject.id,
            'school_class': school_class.id,
            'session': session.id,
            'term': term.id,
        },
        format='multipart',
    )
    assert response.status_code == status.HTTP_400_BAD_REQUEST


@pytest.mark.django_db
def test_student_cannot_access_banks(student_client):
    response = student_client.get('/api/v1/question-banks/')
    assert response.status_code in (status.HTTP_403_FORBIDDEN,)


@pytest.mark.django_db
def test_convert_requires_correct_answers(teacher_client, subject, school_class, session, term, bank):
    QuizQuestion.objects.create(
        madrasah=bank.madrasah, created_by=bank.created_by, subject=subject,
        school_class=school_class, question_type='short_answer',
        question_text='Explain idgham.', correct_answer='', question_bank=bank)

    response = teacher_client.post(f'/api/v1/question-banks/{bank.id}/convert/')
    assert response.status_code == status.HTTP_400_BAD_REQUEST
    assert 'correct answer' in response.data['error']


@pytest.mark.django_db
def test_convert_to_quiz_publishes_and_assigns(
        teacher_client, subject, school_class, session, term, bank):
    for i, q in enumerate([
        {'type': 'mcq', 'text': 'Q1', 'correct': 'A',
         'options': [{'key': 'A', 'text': 'a'}, {'key': 'B', 'text': 'b'}]},
        {'type': 'short_answer', 'text': 'Q2', 'correct': 'lengthen madd', 'options': []},
    ]):
        QuizQuestion.objects.create(
            madrasah=bank.madrasah, created_by=bank.created_by, subject=subject,
            school_class=school_class, question_type=q['type'],
            question_text=q['text'], correct_answer=q['correct'],
            options=q['options'], question_bank=bank)

    response = teacher_client.post(f'/api/v1/question-banks/{bank.id}/convert/')
    assert response.status_code == status.HTTP_200_OK, response.content
    quiz_data = response.data
    assert quiz_data['status'] == 'published'
    assert quiz_data['is_published'] is True
    assert quiz_data['source_bank'] == bank.id
    assert quiz_data['question_count'] == 2

    bank.refresh_from_db()
    assert bank.converted_quiz_id == quiz_data['id']
    assert bank.converted_quiz.assignments.filter(school_class=school_class).exists()

    # Idempotent: second convert returns same quiz
    response2 = teacher_client.post(f'/api/v1/question-banks/{bank.id}/convert/')
    assert response2.status_code == status.HTTP_200_OK
    assert response2.data['id'] == quiz_data['id']


@pytest.mark.django_db
def test_list_groups_banks_and_delete_unlinks(
        teacher_client, bank, student, madrasah):
    from rest_framework.test import APIClient
    response = teacher_client.get('/api/v1/question-banks/')
    assert response.status_code == status.HTTP_200_OK
    assert len(response.data['results']) == 1
    assert response.data['results'][0]['session_name'] == '2025/2026'
    assert response.data['results'][0]['term_number'] == 1

    student_client = APIClient()
    student_client.force_authenticate(user=student)
    resp = student_client.get(f'/api/v1/question-banks/{bank.id}/')
    assert resp.status_code in (status.HTTP_403_FORBIDDEN, status.HTTP_404_NOT_FOUND)

    q = QuizQuestion.objects.create(
        madrasah=bank.madrasah, created_by=bank.created_by, subject=bank.subject,
        school_class=bank.school_class, question_type='mcq',
        question_text='X', correct_answer='A', question_bank=bank)

    response = teacher_client.delete(f'/api/v1/question-banks/{bank.id}/')
    assert response.status_code == status.HTTP_204_NO_CONTENT
    assert not QuestionBank.objects.filter(pk=bank.id).exists()
    q.refresh_from_db()
    assert q.question_bank_id is None
