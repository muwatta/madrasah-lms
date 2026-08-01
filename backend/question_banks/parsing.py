import re

# Question number start: "1.", "1)", "1 -", "١.", "۱)", "(1)", "Q1." etc.
_QUESTION_RE = re.compile(
    r'^\s*(?:\(?\s*(?:\d{1,3}|[٠-٩۰-۹]{1,3})\s*[.\))：:.\-\u2013]+\s*)\s*(.*)$'
)

# MCQ option start: "A.", "B)", "c)", "(A)", "A -", Arabic "أ.", "ب)" etc.
_OPTION_RE = re.compile(
    r'^\s*(?:\(?\s*([A-Za-z]|[أ-ي])\s*[.\))：:.\-\u2013]+\s*)(.*)$'
)

_ARABIC_LETTERS = 'أبجدهوزحطيكلمنسعفصقرشتثخضظغ'
_ARABIC_INDEX = {ch: chr(65 + i) for i, ch in enumerate(_ARABIC_LETTERS[:26])}

_ANSWER_RE = re.compile(
    r'^\s*(?:answer|الإجابة|الإجابه|الجواب|الاجابة|الاجابه|correct|الرد|solution)\s*[:：]\s*(.+)$',
    re.IGNORECASE,
)

_SKIP_RE = re.compile(r'^\s*(?:[-–—_*]+|page\s*\d+|\s*)\s*$', re.IGNORECASE)

_TRUE_FALSE_TEXTS = {'true', 'false', 'صح', 'خطأ', 'صحيح', 'خطا', 'yes', 'no', 'نعم', 'لا'}


def _normalize(text):
    return re.sub(r'\s+', ' ', text).strip()


def _option_key(label):
    label = label.strip()
    if label in _ARABIC_INDEX:
        return _ARABIC_INDEX[label]
    if len(label) == 1 and label.isalpha():
        return label.upper()
    return None


def _detect_question_type(option_texts, options_count):
    if options_count == 2 and all(t.strip().lower() in _TRUE_FALSE_TEXTS for t in option_texts):
        return 'true_false'
    if options_count >= 2:
        return 'mcq'
    return 'short_answer'


def _extract_docx_text(file_obj):
    from docx import Document
    doc = Document(file_obj)
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [_normalize(c.text) for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                lines.append(' | '.join(cells))
    return lines


def _extract_pdf_text(file_obj):
    from pypdf import PdfReader
    reader = PdfReader(file_obj)
    lines = []
    for page in reader.pages:
        text = page.extract_text() or ''
        for raw in text.splitlines():
            line = _normalize(raw)
            if line:
                lines.append(line)
    return lines


def parse_document(file_obj, file_type):
    """Parse a docx/pdf exam document into a list of question dicts.

    Returns: list of {
        'question_type': 'mcq' | 'true_false' | 'short_answer',
        'question_text': str,
        'options': [{'key': 'A', 'text': '...'}, ...],
        'correct_answer': str,
        'explanation': '',
    }
    """
    if file_type == 'docx':
        lines = _extract_docx_text(file_obj)
    else:
        lines = _extract_pdf_text(file_obj)

    blocks = []
    current = None
    for raw in lines:
        line = _normalize(raw)
        if not line or _SKIP_RE.match(line):
            continue
        m = _QUESTION_RE.match(line)
        if m:
            if current:
                blocks.append(current)
            current = {'question_text': [], 'options': [], 'answers': []}
            text = m.group(1).strip()
            if text:
                current['question_text'].append(text)
            continue
        if current is None:
            # Text before the first numbered question is ignored (header/instructions)
            continue

        ans = _ANSWER_RE.match(line)
        if ans:
            current['answers'].append(ans.group(1).strip())
            continue

        om = _OPTION_RE.match(line)
        if om:
            key = _option_key(om.group(1))
            text = om.group(2).strip()
            if key and text:
                current['options'].append({'key': key, 'text': text})
                continue

        current['question_text'].append(line)

    if current:
        blocks.append(current)

    questions = []
    for block in blocks:
        question_text = _normalize(' '.join(block['question_text']))
        if not question_text:
            continue
        option_texts = [o['text'] for o in block['options']]
        q_type = _detect_question_type(option_texts, len(block['options']))
        correct_answer = ''
        if block['answers']:
            answer_text = _normalize(block['answers'][0])
            key = _option_key(answer_text)
            if key:
                correct_answer = key
            elif q_type == 'true_false':
                lower = answer_text.lower()
                correct_answer = 'A' if lower in ('true', 'صحيح', 'صح', 'نعم', 'yes') else 'B'
            else:
                correct_answer = answer_text

        questions.append({
            'question_type': q_type,
            'question_text': question_text,
            'options': block['options'] if q_type in ('mcq', 'true_false') else [],
            'correct_answer': correct_answer,
            'explanation': '',
        })

    return questions
